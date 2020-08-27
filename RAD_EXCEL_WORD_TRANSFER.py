import pandas as pd
import numpy as np
import os
import docx
import openpyxl
import time
from datetime import datetime
from datetime import timedelta
import re

'''判断项目是否需要写入word'''


def WrittenAlterM1(df):
    rows = len(df)
    index = []
    for row in range(rows):
        filedate = datetime.date(df.loc[row, '填写时间'])  # 对填写时期变为年月日形式
        datenow = datetime.date(datetime.now()) - timedelta(2)  # 对目前时间变为年月日形式 可通过  timedelta(1) 调整时间
        if filedate == datenow and df.loc[row, '是否需提交处长'] == '是' and df.loc[row, '解除权限层级'] == '处长权限':
            index.append(row)
    return index


def WrittenAlterM2(df):
    rows = len(df)
    index = []
    for row in range(rows):
        filedate = datetime.date(df.loc[row, '填写时间'])  # 对填写时期变为年月日形式
        datenow = datetime.date(datetime.now()) - timedelta(2)  # 对目前时间变为年月日形式 可通过  timedelta(1) 调整时间
        if filedate == datenow and df.loc[row, '是否需提交处长'] == '是' \
                and df.loc[row, '解除权限层级'] == '副总权限' and df.loc[row, '是否需提交副总'] == '是':
            index.append(row)
    return index


'''生成对应word文本'''


def GenerateDoc(df, targets=[1], name=['报警事项核查模板']):
    doc = docx.Document('报警事项核查模板.docx')

    '''删除文件夹中已经存在的 报警事项核查副总解除权限.docx以及报警事项核查副总解除权限.docx文件'''
    for folders, subfolders, files in os.walk(os.getcwd()):
        for file in files:
            if file.endswith('报警事项核查副总解除权限.docx') or file.endswith('报警事项核查处长解除权限.docx'):
                os.unlink(file)

    '''excel文件内容写入word'''
    for target in targets:
        t1 = doc.add_heading('报警事项:' + df.loc[target, '报警原因'].strip(), 1)
        subtitle1 = doc.add_paragraph(
            '{0}年{1}月{2}日'.format(df.loc[target, '报警日期'].strftime('%Y'),
                                  df.loc[target, '报警日期'].strftime('%m'),
                                  df.loc[target, '报警日期'].strftime('%d'))
            + '   '
            + df.loc[target, '一级分行'],
            'context1')
        subtitle2 = doc.add_paragraph('客户名称：' + df.loc[target, '客户名称'], 'context1')
        subtitle3 = doc.add_paragraph(
            df.loc[target, '业务模块'].strip() + '：' + str(int(df.loc[target, '申报金额'] / 10000)) + '万元   核查人： '
            + df.loc[target, '报警核查人'].strip(), 'context1')
        # 开始针对各个单元格内容进行编辑
        for i in range(len(df.columns) - 41):  # 开始列为 34列+末尾6列不要
            if df.iloc[target, i + 35] != '无内容':
                context = doc.add_paragraph(df.columns[i + 35].strip() + ': ', 'context')
                contents = re.split('\\\\n|\n',
                                    df.iloc[target, i + 35].strip())  # df.iloc[target,i+9].strip().split('\\n')
                context1 = context.add_run(contents[0].strip())
                context1.bold = False
                # 开始针对单元格内部内容进行编辑，对有回车进行分割
                for content in contents[1:]:
                    context2 = doc.add_paragraph(content.strip(), 'context2')

        subtitle4 = doc.add_paragraph('核查结论:' + df.loc[target, '核查结论'].strip(), 'context1')
        doc.save(
            'C:\\Users\\Administrator\\Desktop\\{}{}.docx'.format(df.loc[target, '填写时间'].strftime('%Y-%m-%d'), name[0]))
    return None


os.chdir('C:\\Users\\Administrator\\Desktop\\RAD报警事项核查登记簿汇总')
path = os.getcwd()
rad_data = pd.read_excel('RAD报警事项核查登记簿.xlsm', parse_dates=True, skiprows=1)
rad_data = rad_data.dropna(how='all')
rad_data.loc[:, '报警日期'] = pd.to_datetime(rad_data.loc[:, '报警日期'])
rad_data = rad_data.fillna('无内容')
rad_data = rad_data.drop(rad_data.loc[rad_data['报警日期'] == '无内容', :].index)
rad_data['申报金额'] = rad_data['申报金额'].astype(float)
print(WrittenAlterM1(rad_data))
print(WrittenAlterM2(rad_data))
GenerateDoc(rad_data, WrittenAlterM2(rad_data), ['报警事项核查副总解除权限'])
